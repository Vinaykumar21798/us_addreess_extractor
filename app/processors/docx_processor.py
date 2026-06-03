from docx import Document

from app.exceptions.custom_exceptions import (
    NoExtractableTextException
)


class DOCXProcessor:

    def extract_segments(self, file) -> list[dict]:
        """
        Extract text from a DOCX document including paragraphs, tables, headers, and footers.
        """
        import io

        try:
            file.file.seek(0)
            doc_bytes = file.file.read()
            doc_stream = io.BytesIO(doc_bytes)
            document = Document(doc_stream)
            segments = []

            # 1. Extract from headers
            for idx, section in enumerate(document.sections):
                header = section.header
                if header:
                    for p_idx, p in enumerate(header.paragraphs):
                        if p.text.strip():
                            segments.append({
                                "text": p.text.strip(),
                                "source": f"Section {idx + 1} Header"
                            })
                    for t_idx, table in enumerate(header.tables):
                        for r_idx, row in enumerate(table.rows):
                            for c_idx, cell in enumerate(row.cells):
                                if cell.text.strip():
                                    segments.append({
                                        "text": cell.text.strip(),
                                        "source": f"Section {idx + 1} Header Table {t_idx + 1} Row {r_idx + 1} Cell {c_idx + 1}"
                                    })

            # 2. Extract from paragraphs (main body)
            for p_idx, paragraph in enumerate(document.paragraphs):
                if paragraph.text.strip():
                    segments.append({
                        "text": paragraph.text.strip(),
                        "source": f"Paragraph {p_idx + 1}"
                    })

            # 3. Extract from tables (main body)
            for t_idx, table in enumerate(document.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            segments.append({
                                "text": cell.text.strip(),
                                "source": f"Table {t_idx + 1} Row {r_idx + 1} Cell {c_idx + 1}"
                            })

            # 4. Extract from footers
            for idx, section in enumerate(document.sections):
                footer = section.footer
                if footer:
                    for p_idx, p in enumerate(footer.paragraphs):
                        if p.text.strip():
                            segments.append({
                                "text": p.text.strip(),
                                "source": f"Section {idx + 1} Footer"
                            })
                    for t_idx, table in enumerate(footer.tables):
                        for r_idx, row in enumerate(table.rows):
                            for c_idx, cell in enumerate(row.cells):
                                if cell.text.strip():
                                    segments.append({
                                        "text": cell.text.strip(),
                                        "source": f"Section {idx + 1} Footer Table {t_idx + 1} Row {r_idx + 1} Cell {c_idx + 1}"
                                    })

            if not segments:
                raise NoExtractableTextException(
                    "No readable text found in the DOCX document."
                )

            return segments

        except NoExtractableTextException:
            raise

        except Exception as e:
            raise NoExtractableTextException(
                f"Unable to read DOCX file: {str(e)}"
            )